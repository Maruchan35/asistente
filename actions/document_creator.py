"""document_creator.py — Creador profesional de documentos para JARVIS.
Soporta Word (.docx) con portada, índice, normas APA 7 / Chicago / profesional,
cabeceras, pies de página, estilos académicos y formato completo.
"""
from __future__ import annotations
import json, os, re, traceback
from copy import deepcopy
from datetime import datetime
from pathlib import Path

# ── Resolución de rutas ───────────────────────────────────────────────────────
_FOLDER_MAP = {
    "desktop": "Desktop",     "escritorio": "Desktop",
    "downloads": "Downloads", "descargas": "Downloads",
    "documents": "Documents", "documentos": "Documents",
    "pictures": "Pictures",   "imagenes": "Pictures",   "imágenes": "Pictures",
    "music": "Music",         "musica": "Music",        "música": "Music",
    "videos": "Videos",       "home": "",
}

def _resolve_save_dir(save_path: str) -> Path:
    home = Path(os.path.expanduser("~"))
    if not save_path:
        return home / "Desktop"
    save_path = save_path.replace("~", str(home)).strip()
    p = Path(save_path)
    if p.is_absolute():
        return p.parent if p.suffix in (".docx", ".xlsx", ".txt", ".pdf") else p
    parts = save_path.replace("\\", "/").split("/")
    first = parts[0].strip().lower()
    rest  = "/".join(parts[1:]) if len(parts) > 1 else ""
    if rest and Path(rest).suffix in (".docx", ".xlsx", ".txt", ".pdf"):
        rest = str(Path(rest).parent) if "/" in rest or "\\" in rest else ""
    if first in _FOLDER_MAP:
        base_name = _FOLDER_MAP[first]
        base = (home / base_name) if base_name else home
        return (base / rest) if rest else base
    for root in [home / "Desktop", home / "Downloads", home / "Documents", home]:
        if root.is_dir():
            for folder in os.listdir(root):
                if folder.lower().replace(" ", "_") == first.replace(" ", "_"):
                    base = root / folder
                    return (base / rest) if rest else base
    return home / "Desktop"

# ── Dependencias ──────────────────────────────────────────────────────────────
def _ensure_docx() -> bool:
    try:
        import docx; return True
    except ImportError:
        import subprocess, sys
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "-q"],
                       capture_output=True)
        try:
            import docx; return True
        except ImportError:
            return False

# ── Helpers de formato docx ───────────────────────────────────────────────────
def _cm(v):
    from docx.shared import Cm; return Cm(v)
def _pt(v):
    from docx.shared import Pt; return Pt(v)
def _rgb(r, g, b):
    from docx.shared import RGBColor; return RGBColor(r, g, b)

def _qn(tag):
    from docx.oxml.ns import qn; return qn(tag)

def _el(tag):
    from docx.oxml import OxmlElement; return OxmlElement(tag)

def _set_cell_bg(cell, hex_color: str):
    """Establecer color de fondo de una celda."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = _el("w:shd")
    shd.set(_qn("w:val"),   "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _set_cell_border(cell, color: str = "4472C4", sz: int = 12):
    """Borde izquierdo colored en una celda (efecto callout)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = _el("w:tcBorders")
    for side in ("left",):
        b = _el(f"w:{side}")
        b.set(_qn("w:val"),   "single")
        b.set(_qn("w:sz"),    str(sz))
        b.set(_qn("w:space"), "0")
        b.set(_qn("w:color"), color)
        tcBorders.append(b)
    # Sin bordes en los demás lados
    for side in ("top", "bottom", "right"):
        b = _el(f"w:{side}")
        b.set(_qn("w:val"), "none")
        tcBorders.append(b)
    tcPr.append(tcBorders)

def _add_bottom_border_para(paragraph, color: str = "4472C4", sz: int = 6):
    """Línea inferior bajo un párrafo (estilo heading profesional)."""
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = _el("w:pBdr")
    bot  = _el("w:bottom")
    bot.set(_qn("w:val"),   "single")
    bot.set(_qn("w:sz"),    str(sz))
    bot.set(_qn("w:space"), "1")
    bot.set(_qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def _add_page_number(paragraph, position: str = "center"):
    """Agregar campo PAGE al párrafo del footer."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    align_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
    }
    paragraph.alignment = align_map.get(position, WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run()
    # Build field: PAGE
    fld_begin = _el("w:fldChar"); fld_begin.set(_qn("w:fldCharType"), "begin")
    instr     = _el("w:instrText"); instr.set(_qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld_sep   = _el("w:fldChar"); fld_sep.set(_qn("w:fldCharType"), "separate")
    fld_end   = _el("w:fldChar"); fld_end.set(_qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)

def _add_toc_field(doc):
    """Insertar campo de Tabla de Contenidos de Word (se actualiza al abrir)."""
    para = doc.add_paragraph()
    run  = para.add_run()
    fld_begin = _el("w:fldChar"); fld_begin.set(_qn("w:fldCharType"), "begin")
    fld_begin.set(_qn("w:dirty"), "true")
    instr = _el("w:instrText")
    instr.set(_qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = _el("w:fldChar"); fld_sep.set(_qn("w:fldCharType"), "separate")
    placeholder = _el("w:t"); placeholder.text = "Haz clic aquí y presiona F9 para actualizar la tabla de contenidos."
    fld_end = _el("w:fldChar"); fld_end.set(_qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)
    run.italic = True
    from docx.shared import RGBColor
    run.font.color.rgb = RGBColor(127, 127, 127)
    return para

def _apply_inline(run, text: str):
    """Aplicar negrita/cursiva inline desde markers **text** y _text_."""
    # Ya procesado en _add_inline_para
    run.text = text

def _add_inline_para(doc_or_cell, text: str, style=None, **kwargs):
    """Agregar párrafo con soporte de **bold** e _italic_ inline."""
    if hasattr(doc_or_cell, "add_paragraph"):
        p = doc_or_cell.add_paragraph(style=style) if style else doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell
    # Parser inline simple: **bold**, _italic_, ***bold+italic***
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|_.*?_|\*.*?\*)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = p.add_run()
        if part.startswith("***") and part.endswith("***"):
            run.text = part[3:-3]; run.bold = True; run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]; run.bold = True
        elif (part.startswith("_") and part.endswith("_")) or \
             (part.startswith("*") and part.endswith("*") and not part.startswith("**")):
            run.text = part[1:-1]; run.italic = True
        else:
            run.text = part
    # Aplicar kwargs al párrafo
    if "font_size" in kwargs:
        for r in p.runs: r.font.size = _pt(kwargs["font_size"])
    if "font_name" in kwargs:
        for r in p.runs: r.font.name = kwargs["font_name"]
    if "color" in kwargs:
        for r in p.runs: r.font.color.rgb = kwargs["color"]
    if "bold_all" in kwargs and kwargs["bold_all"]:
        for r in p.runs: r.bold = True
    if "alignment" in kwargs:
        p.alignment = kwargs["alignment"]
    if "space_before" in kwargs:
        p.paragraph_format.space_before = _pt(kwargs["space_before"])
    if "space_after" in kwargs:
        p.paragraph_format.space_after  = _pt(kwargs["space_after"])
    if "line_spacing" in kwargs:
        from docx.shared import Pt
        from docx.enum.text import WD_LINE_SPACING
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing      = kwargs["line_spacing"]
    if "indent" in kwargs:
        p.paragraph_format.first_line_indent = _cm(kwargs["indent"])
    if "left_indent" in kwargs:
        p.paragraph_format.left_indent = _cm(kwargs["left_indent"])
    return p

# ══════════════════════════════════════════════════════════════════════════════
# ESTILOS POR NORMA
# ══════════════════════════════════════════════════════════════════════════════
_NORMS: dict[str, dict] = {
    "apa": {
        "font_body":        "Times New Roman",
        "font_headings":    "Times New Roman",
        "size_body":        12,
        "size_h1":          12,
        "size_h2":          12,
        "size_h3":          12,
        "line_spacing":     2.0,   # APA: doble espacio
        "margin":           2.54,
        "first_indent":     1.27,  # 0.5 pulgada sangría primera línea
        "h1_color":         _rgb(0, 0, 0),
        "h2_color":         _rgb(0, 0, 0),
        "h1_center":        True,
        "h1_bold":          True,
        "h2_center":        False,
        "h2_bold":          True,
        "h3_italic":        True,
        "h3_bold":          True,
        "h1_border":        False,
        "page_num_pos":     "right",
        "para_space_b":     0,
        "para_space_a":     0,
        "h1_space_before":  0,
        "h1_space_after":   0,
        "h2_space_before":  0,
        "h2_space_after":   0,
        "ref_hanging":      1.27,
    },
    "professional": {
        "font_body":        "Calibri",
        "font_headings":    "Calibri Light",
        "size_body":        11,
        "size_h1":          14,    # 14pt — igual que referencia
        "size_h2":          12,
        "size_h3":          11,
        "line_spacing":     1.15,  # interlineado Word por defecto — igual que referencia
        "margin":           2.54,
        "margin_left":      3.0,
        "margin_right":     2.54,
        "first_indent":     0,     # sin sangría primera línea
        "h1_color":         _rgb(31, 73, 125),
        "h2_color":         _rgb(31, 73, 125),
        "h1_center":        False,
        "h1_bold":          True,
        "h2_center":        False,
        "h2_bold":          True,
        "h3_italic":        False,
        "h3_bold":          True,
        "h1_border":        True,
        "h1_border_color":  "2E74B5",
        "page_num_pos":     "center",
        "para_space_b":     0,
        "para_space_a":     8,     # 8pt después de cada párrafo (Word estándar)
        "h1_space_before":  18,
        "h1_space_after":   6,
        "h2_space_before":  12,
        "h2_space_after":   4,
        "ref_hanging":      1.27,
    },
    "academic": {
        "font_body":        "Times New Roman",
        "font_headings":    "Times New Roman",
        "size_body":        12,
        "size_h1":          14,
        "size_h2":          12,
        "size_h3":          12,
        "line_spacing":     1.5,
        "margin":           3.0,
        "margin_right":     2.5,
        "first_indent":     1.0,
        "h1_color":         _rgb(0, 0, 0),
        "h2_color":         _rgb(0, 0, 0),
        "h1_center":        True,
        "h1_bold":          True,
        "h2_center":        False,
        "h2_bold":          True,
        "h3_italic":        True,
        "h3_bold":          False,
        "h1_border":        False,
        "page_num_pos":     "center",
        "para_space_b":     0,
        "para_space_a":     6,
        "h1_space_before":  18,
        "h1_space_after":   6,
        "h2_space_before":  12,
        "h2_space_after":   4,
        "ref_hanging":      1.27,
    },
}
_NORMS["apa7"]     = _NORMS["apa"]
_NORMS["chicago"]  = _NORMS["academic"]
_NORMS["mla"]      = _NORMS["academic"]
_NORMS["report"]   = _NORMS["professional"]


def _apply_document_styles(doc, st: dict, norm_key: str):
    """Aplicar estilos a los estilos nativos de Word (Heading 1/2/3, Normal).
    Esto garantiza TOC funcional y formato consistente en todo el documento."""
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    fn_body = st["font_body"]
    fn_head = st["font_headings"]
    ls      = st["line_spacing"]
    is_apa  = norm_key in ("apa", "apa7")

    def _set_style_ls(style, spacing):
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = spacing

    # ── Normal ────────────────────────────────────────────────────────────────
    try:
        ns = doc.styles["Normal"]
        ns.font.name = fn_body
        ns.font.size = Pt(st["size_body"])
        _set_style_ls(ns, ls)
        ns.paragraph_format.space_before = Pt(st.get("para_space_b", 0))
        ns.paragraph_format.space_after  = Pt(st.get("para_space_a", 8))
        if st.get("first_indent", 0):
            from docx.shared import Cm
            ns.paragraph_format.first_line_indent = Cm(st["first_indent"] / 2.54)
    except Exception:
        pass

    # ── Heading 1 ─────────────────────────────────────────────────────────────
    try:
        h1 = doc.styles["Heading 1"]
        h1.font.name  = fn_head
        h1.font.size  = Pt(st["size_h1"])
        h1.font.color.rgb = st["h1_color"]
        h1.font.bold  = st["h1_bold"]
        h1.font.italic = False
        _set_style_ls(h1, ls if is_apa else 1.15)
        h1.paragraph_format.space_before = Pt(st.get("h1_space_before", 18))
        h1.paragraph_format.space_after  = Pt(st.get("h1_space_after",  6))
        h1.paragraph_format.first_line_indent = None
        h1.paragraph_format.left_indent = None
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        h1.paragraph_format.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                                         if st["h1_center"] else WD_ALIGN_PARAGRAPH.LEFT)
    except Exception:
        pass

    # ── Heading 2 ─────────────────────────────────────────────────────────────
    try:
        h2 = doc.styles["Heading 2"]
        h2.font.name  = fn_head
        h2.font.size  = Pt(st["size_h2"])
        h2.font.color.rgb = st["h2_color"]
        h2.font.bold  = st["h2_bold"]
        h2.font.italic = False
        _set_style_ls(h2, ls if is_apa else 1.15)
        h2.paragraph_format.space_before = Pt(st.get("h2_space_before", 12))
        h2.paragraph_format.space_after  = Pt(st.get("h2_space_after",  4))
        h2.paragraph_format.first_line_indent = None
        h2.paragraph_format.left_indent = None
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        h2.paragraph_format.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                                         if st["h2_center"] else WD_ALIGN_PARAGRAPH.LEFT)
    except Exception:
        pass

    # ── Heading 3 ─────────────────────────────────────────────────────────────
    try:
        h3 = doc.styles["Heading 3"]
        h3.font.name  = fn_head
        h3.font.size  = Pt(st["size_h3"])
        h3.font.color.rgb = st["h2_color"]
        h3.font.bold  = st.get("h3_bold", True)
        h3.font.italic = st.get("h3_italic", False)
        _set_style_ls(h3, ls if is_apa else 1.15)
        h3.paragraph_format.space_before = Pt(st.get("h2_space_before", 12) - 4)
        h3.paragraph_format.space_after  = Pt(st.get("h2_space_after",  4))
        h3.paragraph_format.first_line_indent = None
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception:
        pass

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
def _build_cover(doc, cover: dict, norm_key: str, st: dict):
    """Construir página de portada completa."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared   import Pt, Cm

    section = doc.sections[0]
    section.top_margin    = Cm(st.get("margin", 2.54))
    section.bottom_margin = Cm(st.get("margin", 2.54))
    section.left_margin   = Cm(st.get("margin_left", st.get("margin", 2.54)))
    section.right_margin  = Cm(st.get("margin_right", st.get("margin", 2.54)))

    fn = st["font_body"]

    def centered_para(text, size, bold=False, italic=False, space_before=0, space_after=6, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = fn
        if color:
            run.font.color.rgb = color
        return p

    is_apa = norm_key in ("apa", "apa7")

    if is_apa:
        # APA 7: institución arriba, título centrado en el tercio superior, etc.
        institution = cover.get("institution", "")
        if institution:
            for _ in range(5): doc.add_paragraph()
            centered_para(institution, 12, space_after=0)

        for _ in range(4): doc.add_paragraph()

        title = cover.get("title", "Sin Título")
        centered_para(title, 12, bold=True, space_before=12, space_after=12)

        subtitle = cover.get("subtitle", "")
        if subtitle:
            centered_para(subtitle, 12, space_after=12)

        authors = cover.get("authors", cover.get("author", ""))
        if isinstance(authors, str): authors = [authors]
        if authors:
            centered_para(", ".join(authors), 12, space_after=0)

        course = cover.get("course", "")
        if course:
            centered_para(course, 12, space_after=0)

        professor = cover.get("professor", cover.get("instructor", ""))
        if professor:
            centered_para(professor, 12, space_after=0)

        date = cover.get("date", datetime.now().strftime("%B %Y"))
        centered_para(date, 12, space_after=0)

    else:
        # Estilo profesional / académico
        institution = cover.get("institution", "")
        if institution:
            for _ in range(3): doc.add_paragraph()
            centered_para(institution.upper(), st["size_body"],
                          bold=True, color=st["h1_color"], space_after=4)

        department = cover.get("department", "")
        if department:
            centered_para(department, st["size_body"], space_after=30)

        for _ in range(3): doc.add_paragraph()

        # Línea decorativa superior
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_bottom_border_para(p, color=st.get("h1_border_color", "4472C4"), sz=12)
        doc.add_paragraph()

        title = cover.get("title", "Sin Título")
        centered_para(title.upper(), st["size_h1"] + 4, bold=True,
                      color=st["h1_color"], space_before=6, space_after=6)

        subtitle = cover.get("subtitle", "")
        if subtitle:
            centered_para(subtitle, st["size_h1"], space_after=24)

        # Línea decorativa inferior
        doc.add_paragraph()
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_bottom_border_para(p2, color=st.get("h1_border_color", "4472C4"), sz=12)

        for _ in range(4): doc.add_paragraph()

        label_items = []
        authors = cover.get("authors", cover.get("author", ""))
        if isinstance(authors, str) and authors: authors = [authors]
        if authors:
            label_items.append(("Autor(es):", ", ".join(authors) if isinstance(authors, list) else authors))

        course = cover.get("course", "")
        if course: label_items.append(("Materia:", course))

        professor = cover.get("professor", cover.get("instructor", ""))
        if professor: label_items.append(("Docente:", professor))

        date = cover.get("date", datetime.now().strftime("%B %Y"))
        label_items.append(("Fecha:", date))

        for label, value in label_items:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            r1 = p.add_run(f"{label} ")
            r1.bold = True
            r1.font.name = fn
            r1.font.size = Pt(st["size_body"])
            r2 = p.add_run(value)
            r2.font.name = fn
            r2.font.size = Pt(st["size_body"])

    # Salto de página al final de la portada
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CABECERA Y PIE DE PÁGINA
# ══════════════════════════════════════════════════════════════════════════════
def _setup_header_footer(doc, header_text: str, st: dict, norm_key: str):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    section = doc.sections[-1]
    fn = st["font_body"]

    # Header
    if header_text:
        section.different_first_page_header_footer = True
        hdr = section.header
        hdr_para = hdr.paragraphs[0]
        hdr_para.clear()
        hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if norm_key in ("apa","apa7") else WD_ALIGN_PARAGRAPH.LEFT
        hdr_para.paragraph_format.space_after = Pt(0)
        run = hdr_para.add_run(header_text)
        run.font.name = fn
        run.font.size = _pt(9)
        run.font.color.rgb = _rgb(127, 127, 127)
        # Línea inferior en el header (estilo profesional)
        if norm_key not in ("apa", "apa7"):
            _add_bottom_border_para(hdr_para, color="AAAAAA", sz=4)

    # Footer con número de página
    ftr = section.footer
    ftr_para = ftr.paragraphs[0]
    ftr_para.clear()
    ftr_para.paragraph_format.space_before = Pt(6)
    ftr_para.paragraph_format.space_after  = Pt(0)

    pos = st.get("page_num_pos", "center")
    _add_page_number(ftr_para, pos)

    for r in ftr_para.runs:
        r.font.name = fn
        r.font.size = _pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# PARSER DE CONTENIDO
# ══════════════════════════════════════════════════════════════════════════════
def _parse_and_add_content(doc, content: str, st: dict, norm_key: str, parent=None):
    """Parsear texto estructurado (markdown extendido) y agregar al documento."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    parent = parent if parent is not None else doc

    fn_body = st["font_body"]
    fn_head = st["font_headings"]
    ls      = st["line_spacing"]
    is_apa  = norm_key in ("apa", "apa7")
    fi      = st.get("first_indent", 0)

    lines = content.split("\n")
    i     = 0
    in_list     = False
    list_type   = None
    list_num    = 1

    def _para(text, level=0, is_ref=False):
        nonlocal in_list, list_type, list_num
        in_list = False

        from docx.enum.text import WD_LINE_SPACING
        p = parent.add_paragraph()
        p.paragraph_format.space_before  = _pt(st.get("para_space_b", 0))
        p.paragraph_format.space_after   = _pt(st.get("para_space_a", 6))
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing  = ls

        if is_apa:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if not is_ref:
                p.paragraph_format.first_line_indent = _cm(fi / 2.54) if fi else None
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if is_ref:
            p.paragraph_format.first_line_indent = _cm(0)
            p.paragraph_format.left_indent = _cm(st.get("ref_hanging", 1.27) / 2.54)

        # Inline parsing
        parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|_.*?_)', text)
        for part in parts:
            if not part: continue
            run = p.add_run()
            if part.startswith("***") and part.endswith("***"):
                run.text = part[3:-3]; run.bold = True; run.italic = True
            elif part.startswith("**") and part.endswith("**"):
                run.text = part[2:-2]; run.bold = True
            elif part.startswith("_") and part.endswith("_"):
                run.text = part[1:-1]; run.italic = True
            else:
                run.text = part
            run.font.name = fn_body
            run.font.size = _pt(st["size_body"])
        return p

    def _heading(text, level: int):
        nonlocal in_list, list_type, list_num
        in_list = False

        # Usar estilos nativos de Word → TOC correcto + jerarquía visual real
        style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
        style_name = style_map.get(level, "Heading 3")
        p = parent.add_paragraph(text, style=style_name)

        # Para APA: el texto va en mayúsculas centrado en H1/H2
        if is_apa:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            if level in (1, 2):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Borde inferior solo para professional/report con H1
        if level == 1 and st.get("h1_border"):
            _add_bottom_border_para(p, color=st.get("h1_border_color", "4472C4"))

        return p

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Vacío
        if not stripped:
            i += 1
            in_list = False
            list_num = 1
            continue

        # Salto de página
        if stripped in ("---", "===", "[PAGEBREAK]"):
            if hasattr(parent, "add_page_break"): parent.add_page_break()
            i += 1; continue

        # Tabla de Contenidos
        if stripped.upper() in ("[TOC]", "[INDICE]", "[ÍNDICE]"):
            p = _heading("TABLA DE CONTENIDOS", 1)
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if hasattr(parent, "add_paragraph"): parent.add_paragraph()
            _add_toc_field(parent)
            if hasattr(parent, "add_page_break"): parent.add_page_break()
            i += 1; continue

        # Bloque ABSTRACT
        if stripped.upper() in ("[ABSTRACT]", "[RESUMEN]"):
            _heading("Abstract" if stripped.upper() == "[ABSTRACT]" else "Resumen", 1)
            i += 1
            abs_lines = []
            while i < len(lines) and lines[i].strip().upper() not in ("[/ABSTRACT]", "[/RESUMEN]"):
                abs_lines.append(lines[i])
                i += 1
            _para(" ".join(abs_lines))
            i += 1; continue

        # Bloque BOX / CALLOUT
        m_box = re.match(r'\[BOX(?:[\s]+type="([^"]*)")?(?:[\s]+title="([^"]*)")?\]', stripped, re.IGNORECASE)
        if m_box:
            box_type = (m_box.group(1) or "info").lower()
            box_title = m_box.group(2) or ""
            i += 1
            box_lines = []
            while i < len(lines) and "[/BOX]" not in lines[i].upper():
                box_lines.append(lines[i].strip())
                i += 1
            
            bg_color = "EEF3FB"
            border_color = st.get("h1_border_color", "2E74B5")
            if box_type == "warning":
                bg_color = "FFF4E5"; border_color = "F5B041"
            elif box_type == "success":
                bg_color = "E8F8F5"; border_color = "2ECC71"
            elif box_type in ("danger", "error"):
                bg_color = "FDEDEC"; border_color = "E74C3C"

            # Crear callout box con tabla 1x1
            tbl = parent.add_table(rows=1, cols=1)
            tbl.style = "Table Grid"
            cell = tbl.cell(0, 0)
            _set_cell_bg(cell, bg_color)
            _set_cell_border(cell, color=border_color, sz=18)
            
            if box_title:
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = _pt(4)
                cp.paragraph_format.space_after  = _pt(2)
                r = cp.add_run(box_title)
                r.bold = True; r.font.name = fn_body; r.font.size = _pt(st["size_body"])
            
            # Limpiar párrafo inicial si no hay título, para evitar espacios en blanco extras.
            if not box_title and len(cell.paragraphs) > 0:
                p_elem = cell.paragraphs[0]._element
                p_elem.getparent().remove(p_elem)
                
            _parse_and_add_content(doc, "\n".join(box_lines), st, norm_key, parent=cell)
            
            if hasattr(parent, "add_paragraph"):
                parent.add_paragraph()
            i += 1; continue

        # Bloque TABLE
        if stripped.upper().startswith("[TABLE]") or stripped.upper() == "[TABLE]":
            i += 1
            tbl_rows = []
            while i < len(lines) and "[/TABLE]" not in lines[i].upper():
                row_cells = [c.strip() for c in lines[i].split("|")]
                if any(row_cells):
                    tbl_rows.append(row_cells)
                i += 1
            if tbl_rows:
                ncols = max(len(r) for r in tbl_rows)
                tbl = parent.add_table(rows=len(tbl_rows), cols=ncols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(tbl_rows):
                    for ci, val in enumerate(row):
                        if ci >= ncols: continue
                        cell = tbl.cell(ri, ci)
                        if len(cell.paragraphs) > 0:
                            cell.paragraphs[0].text = val
                            run = cell.paragraphs[0].runs
                        else:
                            p = cell.add_paragraph(val)
                            run = p.runs
                        if ri == 0:  # Header row
                            if run: run[0].bold = True
                            _set_cell_bg(cell, "2E74B5" if norm_key != "apa" else "000000")
                            if run:
                                run[0].font.color.rgb = _rgb(255, 255, 255)
                if hasattr(parent, "add_paragraph"):
                    parent.add_paragraph()
            i += 1; continue

        # Referencias
        if stripped.upper() in ("[REFERENCES]", "[REFERENCIAS]", "[BIBLIOGRAFÍA]", "[BIBLIOGRAFIA]"):
            _heading("Referencias" if "REFER" in stripped.upper() else "Bibliografía", 1)
            i += 1
            while i < len(lines):
                ref_line = lines[i].strip()
                if not ref_line:
                    i += 1; continue
                if ref_line.startswith("[/"):
                    i += 1; break
                # Quitar prefijo [REF] si existe
                if ref_line.upper().startswith("[REF]"):
                    ref_line = ref_line[5:].strip()
                _para(ref_line, is_ref=True)
                i += 1
            continue

        # Línea de referencia suelta
        if stripped.upper().startswith("[REF]"):
            _para(stripped[5:].strip(), is_ref=True)
            i += 1; continue

        # Headings
        if stripped.startswith("#### "):
            _heading(stripped[5:], 3); i += 1; continue
        if stripped.startswith("### "):
            _heading(stripped[4:], 3); i += 1; continue
        if stripped.startswith("## "):
            _heading(stripped[3:], 2); i += 1; continue
        if stripped.startswith("# "):
            _heading(stripped[2:], 1); i += 1; continue

        # Listas con viñetas
        if stripped.startswith("- ") or stripped.startswith("* "):
            from docx.enum.text import WD_LINE_SPACING
            p = parent.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = _pt(2)
            p.paragraph_format.space_after  = _pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = ls
            for part in re.split(r'(\*\*.*?\*\*|_.*?_)', stripped[2:]):
                run = p.add_run()
                if part.startswith("**") and part.endswith("**"):
                    run.text = part[2:-2]; run.bold = True
                elif part.startswith("_") and part.endswith("_"):
                    run.text = part[1:-1]; run.italic = True
                else:
                    run.text = part
                run.font.name = fn_body; run.font.size = _pt(st["size_body"])
            in_list = True; list_type = "bullet"
            i += 1; continue

        # Listas numeradas
        m_num = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m_num:
            from docx.enum.text import WD_LINE_SPACING
            p = parent.add_paragraph(style="List Number")
            p.paragraph_format.space_before = _pt(2)
            p.paragraph_format.space_after  = _pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = ls
            run = p.add_run(m_num.group(2))
            run.font.name = fn_body; run.font.size = _pt(st["size_body"])
            in_list = True; list_type = "number"; list_num += 1
            i += 1; continue

        # Cita en bloque
        if stripped.startswith("> "):
            from docx.enum.text import WD_LINE_SPACING
            p = parent.add_paragraph()
            p.paragraph_format.left_indent  = _cm(1.27)
            p.paragraph_format.right_indent = _cm(1.27)
            p.paragraph_format.space_before = _pt(6)
            p.paragraph_format.space_after  = _pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = ls
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.name = fn_body; run.font.size = _pt(st["size_body"])
            i += 1; continue

        # Párrafo normal
        in_list = False
        _para(stripped)
        i += 1

# ══════════════════════════════════════════════════════════════════════════════
# FUNCIÓN PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════
def document_creator(parameters: dict, player=None) -> str:
    action    = parameters.get("action", "word").lower()
    title     = parameters.get("title", "Documento")
    content   = parameters.get("content", "")
    save_path = parameters.get("save_path", "").strip()
    norm_key  = parameters.get("norm", parameters.get("style", "professional")).lower()
    sheets    = parameters.get("sheets", [])

    # Sobrescritura: si se pasa un path absoluto, usar EXACTAMENTE ese archivo
    # (no agregar timestamp ni crear duplicados). Útil para streaming de deep_research.
    overwrite_existing = parameters.get("overwrite_existing", "")

    # Datos de portada
    cover = parameters.get("cover", None)
    if cover and isinstance(cover, str):
        try: cover = json.loads(cover)
        except Exception: cover = {"title": cover}

    # Encabezado de página
    header_text = parameters.get("header", parameters.get("running_head", ""))
    include_toc = parameters.get("toc", parameters.get("table_of_contents", False))
    if isinstance(include_toc, str):
        include_toc = include_toc.lower() in ("true", "yes", "sí", "si", "1")

    # Resolver directorio
    save_dir = _resolve_save_dir(save_path)
    save_dir.mkdir(parents=True, exist_ok=True)

    timestamp  = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(" ", "_")[:50] or "Documento"

    def log(msg):
        if player: player.write_log(msg)

    try:
        # ── WORD ──────────────────────────────────────────────────────────────
        if action in ("word", "docx", "google_doc", "academic", "apa", "professional"):

            if not _ensure_docx():
                return "Error: No se pudo instalar python-docx. Ejecutá: pip install python-docx"

            from docx import Document
            from docx.enum.text import WD_LINE_SPACING

            doc = Document()

            # Estilo
            if norm_key not in _NORMS:
                norm_key = "professional"
            st = _NORMS[norm_key]

            # Márgenes globales (primera sección)
            sec = doc.sections[0]
            sec.top_margin    = _cm(st.get("margin", 2.54))
            sec.bottom_margin = _cm(st.get("margin", 2.54))
            sec.left_margin   = _cm(st.get("margin_left", st.get("margin", 2.54)))
            sec.right_margin  = _cm(st.get("margin_right", st.get("margin", 2.54)))

            # Aplicar estilos nativos al documento (Normal + Heading 1/2/3)
            _apply_document_styles(doc, st, norm_key)

            # ── PORTADA ──────────────────────────────────────────────────────
            if cover:
                if "title" not in cover: cover["title"] = title
                _build_cover(doc, cover, norm_key, st)
            else:
                # Portada mínima con solo el título
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                for _ in range(8): doc.add_paragraph()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(title)
                run.bold = True
                run.font.size = _pt(st["size_h1"] + 4)
                run.font.name = st["font_headings"]
                run.font.color.rgb = st["h1_color"]
                doc.add_page_break()

            # ── CABECERA + PIE ────────────────────────────────────────────────
            if not header_text and cover:
                header_text = cover.get("course", title)
            _setup_header_footer(doc, header_text or title, st, norm_key)

            # ── TABLA DE CONTENIDOS ───────────────────────────────────────────
            if include_toc and "[TOC]" not in content.upper():
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                # Usar función que maneja estilos nativos
                _parse_and_add_content(doc, "[TOC]", st, norm_key)

            # ── CONTENIDO PRINCIPAL ───────────────────────────────────────────
            if content.strip():
                _parse_and_add_content(doc, content, st, norm_key)

            # Modo override: usar exactamente el path solicitado (sin timestamp)
            if overwrite_existing:
                file_path = Path(overwrite_existing)
                if not file_path.is_absolute():
                    file_path = save_dir / file_path.name
                file_path.parent.mkdir(parents=True, exist_ok=True)
                file_name = file_path.name
            else:
                file_name = f"{safe_title}_{timestamp}.docx"
                file_path = save_dir / file_name
            doc.save(str(file_path))
            log(f"📄 Documento creado: {file_path}")

            info = f"Documento Word profesional creado: '{file_name}' en '{file_path.parent}'."
            if include_toc or "[TOC]" in content.upper():
                info += "\nNOTA: Para actualizar la tabla de contenidos en Word: Ctrl+A y luego F9."
            return info

        # ── EXCEL ─────────────────────────────────────────────────────────────
        elif action in ("excel", "xlsx", "google_sheet"):
            try:
                from openpyxl import Workbook
                from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            except ImportError:
                import subprocess, sys
                subprocess.run([sys.executable, "-m", "pip", "install", "openpyxl", "-q"],
                               capture_output=True)
                from openpyxl import Workbook
                from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

            wb = Workbook()
            wb.remove(wb.active)

            if not sheets:
                ws = wb.create_sheet(title="Hoja1")
                header_fill = PatternFill("solid", fgColor="2E74B5")
                header_font = Font(bold=True, color="FFFFFF", name="Calibri")
                for i, line in enumerate(content.split('\n'), 1):
                    if line.strip():
                        cells = [c.strip() for c in line.split('|')]
                        for j, val in enumerate(cells, 1):
                            cell = ws.cell(row=i, column=j, value=val)
                            cell.font = Font(name="Calibri", size=11)
                            cell.alignment = Alignment(wrap_text=True)
                            if i == 1:  # header row
                                cell.fill = header_fill
                                cell.font = header_font
                # Auto-width
                for col in ws.columns:
                    max_len = max((len(str(c.value or "")) for c in col), default=0)
                    ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
            else:
                for sheet_data in sheets:
                    ws = wb.create_sheet(title=sheet_data.get("name", "Hoja")[:31])
                    headers = sheet_data.get("headers", [])
                    rows    = sheet_data.get("rows", [])
                    if headers:
                        for j, h in enumerate(headers, 1):
                            cell = ws.cell(row=1, column=j, value=h)
                            cell.fill = PatternFill("solid", fgColor="2E74B5")
                            cell.font = Font(bold=True, color="FFFFFF", name="Calibri")
                    for ri, row in enumerate(rows, 2 if headers else 1):
                        for ci, val in enumerate(row, 1):
                            cell = ws.cell(row=ri, column=ci, value=val)
                            cell.font = Font(name="Calibri", size=11)
                            if ri % 2 == 0:
                                cell.fill = PatternFill("solid", fgColor="EEF3FB")
                    for col in ws.columns:
                        max_len = max((len(str(c.value or "")) for c in col), default=0)
                        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)

            file_name = f"{safe_title}_{timestamp}.xlsx"
            file_path = save_dir / file_name
            wb.save(str(file_path))
            log(f"📊 Excel creado: {file_path}")
            return f"Planilla Excel creada: '{file_name}' en '{save_dir}'."

        # ── TEXTO PLANO ────────────────────────────────────────────────────────
        elif action == "text":
            file_name = f"{safe_title}_{timestamp}.txt"
            file_path = save_dir / file_name
            with open(str(file_path), "w", encoding="utf-8") as f:
                f.write(f"{title}\n{'=' * len(title)}\n\n{content}")
            log(f"📝 TXT creado: {file_path}")
            return f"Archivo de texto creado: '{file_name}' en '{save_dir}'."

        else:
            return (f"Acción '{action}' no reconocida. "
                    "Usa: word | excel | text — con norm: apa | apa7 | professional | academic | chicago | mla")

    except Exception as e:
        traceback.print_exc()
        return f"Error al crear el documento: {e}"
